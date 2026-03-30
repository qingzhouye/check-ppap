# -*- coding: utf-8 -*-
"""
SGMW零件一致性核验系统 - 项目立项书文档生成器
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_run_font(run, font_name='宋体', font_size=10.5, bold=False):
    """设置文字格式"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.font.bold = bold

def add_heading_custom(doc, text, level=1):
    """添加自定义标题"""
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    if level == 1:
        set_run_font(run, '黑体', 16, True)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        set_run_font(run, '黑体', 14, True)
    else:
        set_run_font(run, '黑体', 12, True)
    return heading

def add_paragraph_custom(doc, text, font_name='宋体', font_size=10.5, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """添加自定义段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, font_name, font_size, bold)
    para.alignment = alignment
    return para

def create_doc():
    # 创建文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)
    
    # ========== 封面 ==========
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    run = title.add_run('项 目 立 项 书')
    set_run_font(run, '黑体', 26, True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('SGMW零件一致性确认自动化校验系统')
    set_run_font(run, '黑体', 18, True)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 项目信息表格
    info_table = doc.add_table(rows=6, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    
    info_data = [
        ('项目编号', 'SGMW-CC-2024-001'),
        ('项目版本', 'V1.0.1'),
        ('编制日期', '2024年'),
        ('编制部门', '供应商质量监测组'),
        ('文档状态', '正式版'),
        ('密级', '内部使用'),
    ]
    
    for i, (key, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = value
        set_run_font(row.cells[0].paragraphs[0].runs[0], '黑体', 12, True)
        set_run_font(row.cells[1].paragraphs[0].runs[0], '宋体', 12)
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(6)
        set_cell_shading(row.cells[0], 'F0F0F0')
    
    doc.add_page_break()
    
    # ========== 目录 ==========
    add_heading_custom(doc, '目  录', 1)
    doc.add_paragraph()
    
    toc_items = [
        '一、项目背景',
        '二、项目目标',
        '三、项目范围',
        '四、项目组织',
        '五、技术方案',
        '六、校验逻辑对应关系',
        '七、风险分析',
        '八、验收标准',
        '九、项目预算',
        '十、项目审批',
    ]
    
    for item in toc_items:
        para = doc.add_paragraph(item, style='List Number')
        set_run_font(para.runs[0], '宋体', 12)
    
    doc.add_page_break()
    
    # ========== 一、项目背景 ==========
    add_heading_custom(doc, '一、项目背景', 2)
    
    add_heading_custom(doc, '1.1 业务背景', 3)
    add_paragraph_custom(doc, 
        '上汽通用五菱汽车股份有限公司(SGMW)供应商质量管理过程中，需要对供应商提报的零件进行"一致性确认"审核。'
        '该工作涉及核对零件基本信息、验证生产企业名称与供应商一致性、确认CCC认证标识符合性、核对型号规格信息、审核实物图片中的标识信息等。')
    
    add_paragraph_custom(doc, '目前该工作完全依赖人工完成，存在以下问题：')
    problems = [
        '任务量大，人工审核效率低',
        '重复性工作多，容易出错',
        '图片识别需要人工查看，耗时费力',
        '缺乏系统化的校验记录和追溯'
    ]
    for p in problems:
        para = doc.add_paragraph(p, style='List Bullet')
        set_run_font(para.runs[0], '宋体', 10.5)
    
    add_heading_custom(doc, '1.2 现状分析', 3)
    
    # 现状分析表格
    table1 = doc.add_table(rows=5, cols=3)
    table1.style = 'Table Grid'
    headers = ['现状问题', '影响程度', '改进需求']
    for i, h in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = h
        set_run_font(cell.paragraphs[0].runs[0], '黑体', 10.5, True)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    rows_data = [
        ('人工核对Excel关键件清单', '高', '自动比对'),
        ('人工查看CCC证书图片', '高', 'AI自动识别'),
        ('人工查看型号标识图片', '高', 'AI自动识别'),
        ('手动填写审核意见', '中', '自动填写'),
    ]
    for i, (col1, col2, col3) in enumerate(rows_data, 1):
        table1.rows[i].cells[0].text = col1
        table1.rows[i].cells[1].text = col2
        table1.rows[i].cells[2].text = col3
    
    # ========== 二、项目目标 ==========
    doc.add_page_break()
    add_heading_custom(doc, '二、项目目标', 2)
    
    add_heading_custom(doc, '2.1 总体目标', 3)
    add_paragraph_custom(doc, 
        '开发一套浏览器插件系统，实现SGMW零件一致性确认任务的自动化校验和审核，提升工作效率，降低人工错误率。')
    
    add_heading_custom(doc, '2.2 具体目标', 3)
    
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Table Grid'
    headers2 = ['目标项', '目标值', '衡量标准']
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        set_run_font(cell.paragraphs[0].runs[0], '黑体', 10.5, True)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    rows_data2 = [
        ('校验效率提升', '提升80%以上', '单任务处理时间从5分钟降至1分钟以内'),
        ('AI识别准确率', '90%以上', 'CCC标识和型号识别准确率'),
        ('批量处理能力', '单次500+任务', '支持分页获取大量任务'),
        ('数据持久化', '100%', '校验结果和日志自动保存'),
    ]
    for i, (col1, col2, col3) in enumerate(rows_data2, 1):
        table2.rows[i].cells[0].text = col1
        table2.rows[i].cells[1].text = col2
        table2.rows[i].cells[2].text = col3
    
    # ========== 三、项目范围 ==========
    add_heading_custom(doc, '三、项目范围', 2)
    
    add_heading_custom(doc, '3.1 功能范围', 3)
    add_paragraph_custom(doc, '【核心功能】', bold=True)
    core_features = [
        '任务列表自动提取',
        '关键件Excel清单导入管理',
        '智能一致性校验（7大校验项）',
        'AI图像识别（CCC标识、型号标识）',
        '批量校验处理',
        '批量自动审核提交',
        '人工审核干预',
        '操作日志和结果导出'
    ]
    for f in core_features:
        para = doc.add_paragraph(f, style='List Bullet')
        set_run_font(para.runs[0], '宋体', 10.5)
    
    add_paragraph_custom(doc, '【辅助功能】', bold=True)
    aux_features = [
        '窗口自由拖动定位',
        '数据持久化存储',
        '多Excel文件管理',
        'API Key配置管理',
        '进度可视化展示'
    ]
    for f in aux_features:
        para = doc.add_paragraph(f, style='List Bullet')
        set_run_font(para.runs[0], '宋体', 10.5)
    
    add_heading_custom(doc, '3.2 技术范围', 3)
    table3 = doc.add_table(rows=5, cols=2)
    table3.style = 'Table Grid'
    headers3 = ['技术项', '技术选型']
    for i, h in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = h
        set_run_font(cell.paragraphs[0].runs[0], '黑体', 10.5, True)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    tech_data = [
        ('浏览器插件架构', 'Manifest V3'),
        ('前端技术', '原生JavaScript + CSS'),
        ('Excel处理', 'SheetJS (xlsx.full.min.js)'),
        ('AI识别服务', '智谱AI GLM-4V-Flash'),
    ]
    for i, (col1, col2) in enumerate(tech_data, 1):
        table3.rows[i].cells[0].text = col1
        table3.rows[i].cells[1].text = col2
    
    # ========== 四、项目组织 ==========
    add_heading_custom(doc, '四、项目组织', 2)
    
    add_heading_custom(doc, '4.1 项目团队', 3)
    table4 = doc.add_table(rows=6, cols=3)
    table4.style = 'Table Grid'
    headers4 = ['角色', '职责', '人员']
    for i, h in enumerate(headers4):
        cell = table4.rows[0].cells[i]
        cell.text = h
        set_run_font(cell.paragraphs[0].runs[0], '黑体', 10.5, True)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    team_data = [
        ('项目发起人', '项目立项、资源协调、验收', '部门负责人'),
        ('产品经理', '需求分析、功能设计、用户验收', '(待指定)'),
        ('技术负责人', '技术架构、代码审查、技术决策', '(待指定)'),
        ('前端开发', '插件开发、页面交互实现', '(待指定)'),
        ('测试人员', '功能测试、兼容性测试、性能测试', '(待指定)'),
    ]
    for i, (col1, col2, col3) in enumerate(team_data, 1):
        table4.rows[i].cells[0].text = col1
        table4.rows[i].cells[1].text = col2
        table4.rows[i].cells[2].text = col3
    
    add_heading_custom(doc, '4.2 项目里程碑', 3)
    table5 = doc.add_table(rows=6, cols=4)
    table5.style = 'Table Grid'
    headers5 = ['阶段', '主要工作', '计划周期', '交付物']
    for i, h in enumerate(headers5):
        cell = table5.rows[0].cells[i]
        cell.text = h
        set_run_font(cell.paragraphs[0].runs[0], '黑体', 10.5, True)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    milestone_data = [
        ('需求分析', '需求调研、功能设计', '1周', '需求文档'),
        ('技术设计', '架构设计、接口设计', '1周', '设计文档'),
        ('开发实现', '编码开发、单元测试', '3周', '源代码'),
        ('测试验证', '功能测试、Bug修复', '2周', '测试报告'),
        ('上线部署', '用户培训、正式发布', '1周', '上线版本'),
    ]
    for i, (col1, col2, col3, col4) in enumerate(milestone_data, 1):
        table5.rows[i].cells[0].text = col1
        table5.rows[i].cells[1].text = col2
        table5.rows[i].cells[2].text = col3
        table5.rows[i].cells[3].text = col4
    
    # ========== 五、技术方案 ==========
    doc.add_page_break()
    add_heading_custom(doc, '五、技术方案', 2)
    
    add_heading_custom(doc, '5.1 系统架构', 3)
    add_paragraph_custom(doc, 
        '系统采用浏览器扩展架构，包含Popup弹窗层、Content脚本层、Background后台服务层三层结构，'
        '通过Chrome Storage API实现数据持久化，并集成SGMW业务系统API和智谱AI图像识别服务。')
    
    add_heading_custom(doc, '5.2 核心模块设计', 3)
    modules = [
        ('任务提取模块', '支持API直接获取任务列表、DOM解析提取任务列表、layui table cache读取'),
        ('数据校验模块', '7大校验项并行处理，支持API数据直接校验和弹窗页面数据校验'),
        ('AI识别模块', '图片Base64编码传输，CCC标识和型号标识专用Prompt，识别结果JSON解析'),
        ('批量处理模块', '任务队列管理、自动页面操作、结果持久化存储、进度实时反馈'),
        ('人工审核模块', '结果可视化展示、人工确认/拒绝操作、审核备注记录、状态同步更新'),
    ]
    for name, desc in modules:
        para = doc.add_paragraph()
        run = para.add_run(f'【{name}】')
        set_run_font(run, '黑体', 10.5, True)
        run = para.add_run(f' {desc}')
        set_run_font(run, '宋体', 10.5)
    
    add_heading_custom(doc, '5.3 接口清单', 3)
    table6 = doc.add_table(rows=9, cols=3)
    table6.style = 'Table Grid'
    headers6 = ['接口类型', '接口名称', '用途']
    for i, h in enumerate(headers6):
        cell = table6.rows[0].cells[i]
        cell.text = h
        set_run_font(cell.paragraphs[0].runs[0], '黑体', 10.5, True)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    api_data = [
        ('内部API', 'GET_EXCEL_LIST', '获取已导入Excel列表'),
        ('内部API', 'ADD_EXCEL_DATA', '添加Excel数据'),
        ('内部API', 'QUERY_PART', '查询零件信息'),
        ('内部API', 'RECOGNIZE_IMAGE', 'AI识别图片'),
        ('内部API', 'SAVE_BATCH_RESULTS', '保存批量校验结果'),
        ('外部API', 'listUniCheckTaskSearch', '获取任务列表'),
        ('外部API', 'getUniCheckTaskInfo', '获取任务详情'),
        ('外部API', 'chat/completions', 'AI图像识别'),
    ]
    for i, (col1, col2, col3) in enumerate(api_data, 1):
        table6.rows[i].cells[0].text = col1
        table6.rows[i].cells[1].text = col2
        table6.rows[i].cells[2].text = col3
    
    # ========== 六、校验逻辑对应关系 ==========
    doc.add_page_break()
    add_heading_custom(doc, '六、校验逻辑对应关系', 2)
    
    add_heading_custom(doc, '6.1 校验流程说明', 3)
    add_paragraph_custom(doc, 
        '系统校验流程分为两种方式：API方式（直接请求后端接口获取数据）和弹窗方式（自动点击打开详情弹窗获取数据）。'
        '优先使用API方式，失败时自动降级到弹窗方式。校验完成后根据结果判定为通过、需人工或不通过。')
    
    add_heading_custom(doc, '6.2 校验字段对应关系表', 3)
    table7 = doc.add_table(rows=8, cols=5)
    table7.style = 'Table Grid'
    headers7 = ['序号', '校验项目', '页面数据来源', 'Excel数据来源', '比对规则']
    for i, h in enumerate(headers7):
        cell = table7.rows[0].cells[i]
        cell.text = h
        set_run_font(cell.paragraphs[0].runs[0], '黑体', 9, True)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    check_data = [
        ('1', '基本信息', 'API: partsName, latestPartsCode, carType', '-', '仅展示'),
        ('2', '生产企业名称一致性', 'API: manufacturerList中的manufacturerName', '-', '供应商与生产企业名称匹配'),
        ('3', '是否CCC件', 'API: isCccParts', 'Excel: ccc列', '布尔值比对'),
        ('4', '型号信息', 'API: modelList中的model', 'Excel: modelSpec列', '字符串包含匹配'),
        ('5', '生产企业(与Excel)', 'API: manufacturerList', 'Excel: manufacturer列', '字符串包含匹配'),
        ('6', 'CCC标识(AI识别)', '附件: cccFile', '-', 'AI识别图片中CCC标志'),
        ('7', '型号标识(AI识别)', '附件: modelFile', '-', 'AI识别图片中型号代码'),
    ]
    for i, (col1, col2, col3, col4, col5) in enumerate(check_data, 1):
        table7.rows[i].cells[0].text = col1
        table7.rows[i].cells[1].text = col2
        table7.rows[i].cells[2].text = col3
        table7.rows[i].cells[3].text = col4
        table7.rows[i].cells[4].text = col5
    
    add_heading_custom(doc, '6.3 Excel文件字段映射', 3)
    table8 = doc.add_table(rows=16, cols=3)
    table8.style = 'Table Grid'
    headers8 = ['Excel列', '字段名', '用途']
    for i, h in enumerate(headers8):
        cell = table8.rows[0].cells[i]
        cell.text = h
        set_run_font(cell.paragraphs[0].runs[0], '黑体', 10.5, True)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    excel_data = [
        ('A列', 'category', '零件类别'),
        ('B列', 'partComponentName', '零件总成名称'),
        ('C列', 'gonggao', '公告'),
        ('D列', 'huanbao', '环保'),
        ('E列', 'ccc', 'CCC标识(●表示是)'),
        ('F列', 'cccCertificate', 'CCC证书编号'),
        ('G列', 'modelSpec', '型号规格'),
        ('H列', 'manufacturer', '生产企业'),
        ('I列', 'certNumber', '认证证书号'),
        ('J列', 'applicableModel', '适用车型'),
        ('N列', 'chinesePartName', '中文零件名称'),
        ('O列', 'partNumber', '零件号'),
        ('P列', 'englishPartName', '英文零件名称'),
    ]
    for i, (col1, col2, col3) in enumerate(excel_data, 1):
        table8.rows[i].cells[0].text = col1
        table8.rows[i].cells[1].text = col2
        table8.rows[i].cells[2].text = col3
    
    add_heading_custom(doc, '6.4 API接口对应关系', 3)
    table9 = doc.add_table(rows=4, cols=4)
    table9.style = 'Table Grid'
    headers9 = ['功能', 'API端点', '方法', '请求参数']
    for i, h in enumerate(headers9):
        cell = table9.rows[0].cells[i]
        cell.text = h
        set_run_font(cell.paragraphs[0].runs[0], '黑体', 10.5, True)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    api_detail_data = [
        ('获取任务列表', '/api/unifomity/uniformityCheckSWTaskSearch/listUniCheckTaskSearch', 'GET', 'pageSize, pageNum'),
        ('获取任务详情', '/api/unifomity/uniformityCheckSWTaskSearch/getUniCheckTaskInfo', 'POST', '{ id: taskId }'),
        ('AI图像识别', 'https://open.bigmodel.cn/api/paas/v4/chat/completions', 'POST', '图片Base64 + Prompt'),
    ]
    for i, (col1, col2, col3, col4) in enumerate(api_detail_data, 1):
        table9.rows[i].cells[0].text = col1
        table9.rows[i].cells[1].text = col2
        table9.rows[i].cells[2].text = col3
        table9.rows[i].cells[3].text = col4
    
    # ========== 七、风险分析 ==========
    add_heading_custom(doc, '七、风险分析', 2)
    
    table10 = doc.add_table(rows=6, cols=3)
    table10.style = 'Table Grid'
    headers10 = ['风险项', '风险等级', '应对措施']
    for i, h in enumerate(headers10):
        cell = table10.rows[0].cells[i]
        cell.text = h
        set_run_font(cell.paragraphs[0].runs[0], '黑体', 10.5, True)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    risk_data = [
        ('页面结构变更', '高', '设计多策略降级方案，API优先+DOM解析备份'),
        ('AI识别准确率不足', '中', '设置人工确认机制，重要决策人工复核'),
        ('API接口变更', '中', '封装API调用层，便于快速适配'),
        ('浏览器兼容性问题', '低', '测试覆盖Chrome/Edge/Firefox主流版本'),
        ('数据安全问题', '中', '数据本地存储，不上传敏感信息'),
    ]
    for i, (col1, col2, col3) in enumerate(risk_data, 1):
        table10.rows[i].cells[0].text = col1
        table10.rows[i].cells[1].text = col2
        table10.rows[i].cells[2].text = col3
    
    # ========== 八、验收标准 ==========
    add_heading_custom(doc, '八、验收标准', 2)
    
    add_heading_custom(doc, '8.1 功能验收标准', 3)
    func_std = [
        '任务列表提取成功率 ≥ 95%',
        'Excel导入解析成功率 ≥ 99%',
        '一致性校验准确率 ≥ 90%',
        'AI图像识别可用率 ≥ 95%',
        '批量处理稳定性 ≥ 99%',
        '数据持久化完整性 100%'
    ]
    for s in func_std:
        para = doc.add_paragraph(s, style='List Bullet')
        set_run_font(para.runs[0], '宋体', 10.5)
    
    add_heading_custom(doc, '8.2 性能验收标准', 3)
    perf_std = [
        '单任务校验时间 ≤ 30秒（含AI识别）',
        '批量处理支持 ≥ 500条任务',
        '页面响应时间 ≤ 2秒',
        '内存占用 ≤ 200MB'
    ]
    for s in perf_std:
        para = doc.add_paragraph(s, style='List Bullet')
        set_run_font(para.runs[0], '宋体', 10.5)
    
    # ========== 九、项目预算 ==========
    add_heading_custom(doc, '九、项目预算', 2)
    
    table11 = doc.add_table(rows=5, cols=3)
    table11.style = 'Table Grid'
    headers11 = ['预算项', '金额(元)', '说明']
    for i, h in enumerate(headers11):
        cell = table11.rows[0].cells[i]
        cell.text = h
        set_run_font(cell.paragraphs[0].runs[0], '黑体', 10.5, True)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    budget_data = [
        ('人力成本', '(内部)', '开发人员、测试人员工时'),
        ('AI服务费用', '(按量计费)', '智谱AI API调用费用'),
        ('测试设备', '(已有)', '开发测试用电脑'),
        ('其他费用', '-', ''),
    ]
    for i, (col1, col2, col3) in enumerate(budget_data, 1):
        table11.rows[i].cells[0].text = col1
        table11.rows[i].cells[1].text = col2
        table11.rows[i].cells[2].text = col3
    
    # ========== 十、项目审批 ==========
    doc.add_page_break()
    add_heading_custom(doc, '十、项目审批', 2)
    
    table12 = doc.add_table(rows=5, cols=4)
    table12.style = 'Table Grid'
    headers12 = ['审批环节', '签字', '日期', '意见']
    for i, h in enumerate(headers12):
        cell = table12.rows[0].cells[i]
        cell.text = h
        set_run_font(cell.paragraphs[0].runs[0], '黑体', 10.5, True)
        set_cell_shading(cell, '4472C4')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    approve_data = [
        ('部门负责人', '', '', ''),
        ('技术负责人', '', '', ''),
        ('质量负责人', '', '', ''),
        ('项目发起人', '', '', ''),
    ]
    for i, (col1, col2, col3, col4) in enumerate(approve_data, 1):
        table12.rows[i].cells[0].text = col1
        table12.rows[i].cells[1].text = col2
        table12.rows[i].cells[2].text = col3
        table12.rows[i].cells[3].text = col4
        # 设置行高以便签字
        table12.rows[i].height = Cm(1.5)
    
    # 保存文档
    output_path = r'c:\Users\Administrator\Desktop\AI-一致性确认 2.0\SGMW零件一致性核验系统_项目立项书.docx'
    doc.save(output_path)
    print(f'文档已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    create_doc()
